"""Word (.docx) resume export — matching the user's exact Chinese resume format.

Produces a document with:
  - Name (14pt bold, centered)
  - Contact line (10pt bold, centered)
  - 教育背景 / 学术经历 / 实习经历 / 竞赛获奖 / 技能其他 sections
  - Title + date on the same line (date flush-right via tab-stop)
  - Bullet descriptions with 两端对齐
  - Bottom skill/language/interest lines (9pt)
"""

from __future__ import annotations

from typing import Any, List, Optional

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Constants matching the reference document ──────────────────────────

FONT_NAME = "微软雅黑"
FONT_NAME_ASCII = "Microsoft YaHei"

SIZE_NAME     = Pt(14)   # 姓名
SIZE_CONTACT  = Pt(10)   # 联系方式
SIZE_SECTION  = Pt(11)   # 教育背景 / 学术经历 etc.
SIZE_TITLE    = Pt(10)   # 学校/项目/公司名称行
SIZE_BULLET   = Pt(10)   # 要点描述
SIZE_BOTTOM   = Pt(9)    # 技能/语言/兴趣

COLOR_BLACK = RGBColor(0, 0, 0)


# ── Helpers ────────────────────────────────────────────────────────────


def _set_font(run, name: str = FONT_NAME, size: Pt = SIZE_BULLET,
              bold: bool = False, color: RGBColor = COLOR_BLACK):
    """Apply font settings to a run, including East-Asian font."""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    # Set East-Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def _add_paragraph(doc, text: str = "", alignment=None,
                   space_after: Pt = Pt(0), space_before: Pt = Pt(0),
                   left_indent: Cm = None):
    """Add a paragraph with spacing and optional alignment."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    if left_indent is not None:
        pf.left_indent = left_indent
    return p


def _add_run(p, text: str, name: str = FONT_NAME, size: Pt = SIZE_BULLET,
             bold: bool = False, color: RGBColor = COLOR_BLACK):
    """Add a run with font settings."""
    run = p.add_run(text)
    _set_font(run, name=name, size=size, bold=bold, color=color)
    return run


def _add_title_line(doc, title: str, date_str: str = ""):
    """Add a bold title line with optional date flush-right via tab stop.

    Uses a right-aligned tab stop at page width for the date.
    """
    p = _add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    # Add right tab stop at page width
    pf = p.paragraph_format
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Cm(16.0))  # right-aligned tab near right margin

    _add_run(p, title, size=SIZE_TITLE, bold=True)
    if date_str:
        _add_run(p, "\t", size=SIZE_TITLE, bold=True)
        _add_run(p, date_str, size=SIZE_TITLE, bold=True)
    return p


def _add_bullet(doc, text: str, bold: bool = False):
    """Add a bullet-point description paragraph (List Bullet style)."""
    p = _add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.74)
    pf.first_line_indent = Cm(-0.37)
    _add_run(p, text, size=SIZE_BULLET, bold=bold)
    return p


def _add_section_header(doc, title: str):
    """Add a section header like '教育背景'."""
    _add_paragraph(doc, space_before=Pt(4))
    p = _add_paragraph(doc, space_after=Pt(2))
    _add_run(p, title, size=SIZE_SECTION, bold=True)
    return p


def _add_bottom_line(doc, label: str, value: str):
    """Add a bottom line like '语言：英语（CET-4 617）'."""
    p = _add_paragraph(doc, space_after=Pt(1))
    _add_run(p, label, size=SIZE_BOTTOM, bold=True)
    _add_run(p, value, size=SIZE_BOTTOM, bold=False)
    return p


def _extract_date_range(start: str, end: str) -> str:
    """Build a date range string like '2023.09 - 2027.06'."""
    s = (start or "").strip().replace("-", ".")
    e = (end or "").strip().replace("-", ".")
    if s and e:
        return f"{s} - {e}"
    elif s:
        return s
    elif e:
        return e
    return ""


# ── Main export function ───────────────────────────────────────────────


def build_resume_docx(
    options: dict,
    student=None,
    courses=None,
    overview=None,
    experiences=None,
    achievements=None,
    roles=None,
    output_path: str = None,
) -> Document:
    """Build a .docx resume matching the reference Chinese resume format.

    Parameters
    ----------
    options : dict
        Form values (name, email, phone, school, major, degree, skills_body, etc.)
    student : Student model (fallback)
    courses : list[Course]
    overview : dict (GPA, weighted_average, etc.)
    experiences : list[Experience]
    achievements : list[Achievement]
    roles : list[Role]
    output_path : str, optional — if provided, save to this path

    Returns
    -------
    docx.Document — the built document (also saved if output_path given)
    """
    experiences = experiences or []
    courses = courses or []
    overview = overview or {}
    achievements = achievements or []

    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = SIZE_BULLET
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # ── Collect data ─────────────────────────────────────────────────
    name = options.get("name") or (student.name if student else "姓名")
    email = options.get("email") or (student.email if student else "")
    phone = options.get("phone") or (student.phone if student else "")
    school = options.get("school") or (student.college if student else "")
    major = options.get("major") or (student.major if student else "")
    degree = options.get("degree") or "本科"

    # Build enrollment year range
    enrollment = (student.enrollment_year or "") if student else ""
    if enrollment and degree:
        start_year = enrollment
        end_year = str(int(enrollment) + 4) if degree == "本科" else str(int(enrollment) + 3)
        edu_dates = f"{start_year}.09 - {end_year}.06"
    else:
        edu_dates = ""

    # ── 1. Name ──────────────────────────────────────────────────────
    p = _add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(0))
    _add_run(p, name, size=SIZE_NAME, bold=True)

    # ── 2. Contact ───────────────────────────────────────────────────
    contact_parts = []
    if phone:
        contact_parts.append(f"手机：(+86) {phone}")
    if email:
        contact_parts.append(f"邮箱：{email}")
    contact_line = "    ".join(contact_parts)
    p = _add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    _add_run(p, contact_line, size=SIZE_CONTACT, bold=True)

    # ── 3. 教育背景 ─────────────────────────────────────────────────
    _add_section_header(doc, "教育背景")

    edu_title = f"{school} | {major} {degree}"
    _add_title_line(doc, edu_title, edu_dates)

    # Core courses
    if courses:
        top_courses = sorted(courses, key=lambda c: c.grade if c.grade else 0, reverse=True)[:10]
        course_text = "核心课程：" + "、".join(
            f"{c.name}（{c.grade:g}）" if c.grade else c.name
            for c in top_courses
        )
        _add_bullet(doc, course_text)

    # GPA
    if overview:
        gpa_text = (
            f"GPA：{overview.get('weighted_average', 0):.2f}/100"
            f" | 专业排名：前{max(1, int((1 - overview.get('gpa', 0)/4.0) * 100))}%"
        )
        _add_bullet(doc, gpa_text)

    # ── 4. 学术经历 ─────────────────────────────────────────────────
    academic_exps = [e for e in experiences if e.exp_type == "学术经历"]
    if academic_exps:
        _add_section_header(doc, "学术经历")

        for exp in academic_exps:
            title_line = f"{exp.title} | {exp.role or ''}"
            if exp.organization:
                title_line += f" | {exp.organization}"
            dates = _extract_date_range(exp.start_date, exp.end_date)
            _add_title_line(doc, title_line, dates)

            if exp.description:
                for line in exp.description.strip().split("\n"):
                    line = line.strip()
                    if line:
                        _add_bullet(doc, line)
            if exp.outcome:
                _add_bullet(doc, exp.outcome)

    # ── 5. 研究经历 ─────────────────────────────────────────────────
    research_exps = [e for e in experiences if e.exp_type == "研究经历"]
    if research_exps:
        _add_section_header(doc, "研究经历")

        for exp in research_exps:
            title_line = f"{exp.title} | {exp.role or ''}"
            if exp.organization:
                title_line += f" | {exp.organization}"
            dates = _extract_date_range(exp.start_date, exp.end_date)
            _add_title_line(doc, title_line, dates)

            if exp.description:
                for line in exp.description.strip().split("\n"):
                    line = line.strip()
                    if line:
                        _add_bullet(doc, line)
            if exp.outcome:
                _add_bullet(doc, exp.outcome)

    # ── 6. 实习经历 ─────────────────────────────────────────────────
    intern_exps = [e for e in experiences if e.exp_type == "实习经历"]
    if intern_exps:
        _add_section_header(doc, "实习经历")

        for exp in intern_exps:
            org = exp.organization or ""
            role = exp.role or ""
            title_line = f"{org}{' | ' if org and role else ''}{role}"
            dates = _extract_date_range(exp.start_date, exp.end_date)
            _add_title_line(doc, title_line, dates)

            if exp.description:
                for line in exp.description.strip().split("\n"):
                    line = line.strip()
                    if line:
                        _add_bullet(doc, line)

    # ── 7. 竞赛经历 ─────────────────────────────────────────────────
    competition_exps = [e for e in experiences if e.exp_type == "竞赛经历"]
    if competition_exps:
        _add_section_header(doc, "竞赛经历")

        for exp in competition_exps:
            title_line = f"{exp.title} | {exp.role or ''}"
            if exp.organization:
                title_line += f" | {exp.organization}"
            dates = _extract_date_range(exp.start_date, exp.end_date)
            _add_title_line(doc, title_line, dates)

            if exp.description:
                for line in exp.description.strip().split("\n"):
                    line = line.strip()
                    if line:
                        _add_bullet(doc, line)
            if exp.outcome:
                _add_bullet(doc, exp.outcome)

    # ── 8. 竞赛获奖（成就）───────────────────────────────────────────
    if achievements:
        _add_section_header(doc, "竞赛获奖")

        for ach in achievements:
            title_line = f"{ach.title} | {ach.ach_type}"
            if ach.issuer:
                title_line += f" | {ach.issuer}"
            _add_title_line(doc, title_line, ach.date or "")

            if ach.description:
                for line in ach.description.strip().split("\n"):
                    line = line.strip()
                    if line:
                        _add_bullet(doc, line)

    # ── 9. 其它 ─────────────────────────────────────────────────────
    other_exps = [e for e in experiences if e.exp_type == "其它"]
    if other_exps:
        _add_section_header(doc, "其它经历")
        for exp in other_exps:
            title_line = f"{exp.title} | {exp.role or ''}"
            if exp.organization:
                title_line += f" | {exp.organization}"
            dates = _extract_date_range(exp.start_date, exp.end_date)
            _add_title_line(doc, title_line, dates)
            if exp.description:
                for line in exp.description.strip().split("\n"):
                    line = line.strip()
                    if line:
                        _add_bullet(doc, line)

    # ── 10. 技能/其他 ────────────────────────────────────────────────
    _add_section_header(doc, "其他")

    # Skills
    skills = options.get("skills_body") or (student.skills if student else "")
    if skills:
        # Flatten: join all lines with commas
        flat_skills = skills.replace("\n", "、").replace("• ", "").replace("  ", " ")
        _add_bottom_line(doc, "技能：", flat_skills)

    # Languages (hardcoded but can be made configurable)
    _add_bottom_line(doc, "语言：", "英语（CET-4/6）")

    # Interests
    _add_bottom_line(doc, "兴趣：", "数据科学、人工智能")

    # Student org
    if roles:
        role_text = "、".join(f"{r.title}（{r.organization}）" for r in roles[:3])
        _add_bottom_line(doc, "学生工作：", role_text)

    # ── Save ─────────────────────────────────────────────────────────
    if output_path:
        doc.save(output_path)

    return doc
